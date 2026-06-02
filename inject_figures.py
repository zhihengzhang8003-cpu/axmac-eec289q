"""
inject_figures.py -- insert 4 figures + update Abstract/Conclusion in v2 report.
Produces Bias_Aware_Approximate_MAC_IEEE_v3.docx
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

SRC = "E:/AIDev/EEC 289Q 002 SQ 2026：Deep Learning Hardware/project/Bias_Aware_Approximate_MAC_IEEE_v2.docx"
DST = "E:/AIDev/EEC 289Q 002 SQ 2026：Deep Learning Hardware/project/Bias_Aware_Approximate_MAC_IEEE_v3.docx"
FIGDIR = "E:/AIDev/EEC 289Q 002 SQ 2026：Deep Learning Hardware/project/experiments/results/figures"

BLUE = RGBColor(0x00, 0x44, 0xCC)

doc = Document(SRC)

def blue_run(para, text, bold=False, size=9):
    run = para.add_run(text)
    run.font.color.rgb = BLUE
    run.font.bold = bold
    run.font.size = Pt(size)
    return run

def insert_before(ref_para, new_elem):
    ref_para._p.addprevious(new_elem)

def new_para_elem(doc, style="Normal"):
    p = doc.add_paragraph()
    p.style = doc.styles[style]
    elem = p._p
    p._p.getparent().remove(p._p)
    return elem, p

# ─── locate key paragraphs ────────────────────────────────────────────────────
paras = doc.paragraphs
def find(text):
    return next(p for p in paras if text in p.text)

vi_para   = find("VI. LIMITATIONS")
concl_para = find("VII. CONCLUSION")

# ─── insert figures before VI. LIMITATIONS ────────────────────────────────────
figs = [
    ("fig1_bias_accumulation.png",
     "Fig. 1. Accumulated truncation bias vs. depth N (INT8, K=6). "
     "trunc grows linearly (∝N); stochastic is bounded by √N."),
    ("fig2_three_way_rounding.png",
     "Fig. 2. Per-MAC accumulated bias for trunc/round/stochastic (INT8). "
     "round nearly eliminates coherent bias at near-zero hardware cost."),
    ("fig3_ppa_hardware.png",
     "Fig. 3. Hardware synthesis results on EP4CE10 (Quartus Prime Lite 18.1). "
     "round: +46 LE vs trunc; stochastic LFSR exceeds the exact baseline."),
    ("fig4_accuracy_vs_K.png",
     "Fig. 4. Top-1 agreement with exact inference vs K (N=1000 INT8 samples). "
     "Hardware board measurements annotated. "
     "At K=6: trunc 74.1%, round 96.8%, stochastic 96.0%."),
]

for fname, caption in reversed(figs):
    fpath = os.path.join(FIGDIR, fname)

    # caption paragraph
    cap_p = doc.add_paragraph()
    cap_p.style = doc.styles["Normal"]
    blue_run(cap_p, caption, bold=False, size=8)
    cap_elem = cap_p._p
    cap_p._p.getparent().remove(cap_p._p)

    # image paragraph
    img_p = doc.add_paragraph()
    img_p.style = doc.styles["Normal"]
    img_p.paragraph_format.space_after = Pt(2)
    run = img_p.add_run()
    run.add_picture(fpath, width=Inches(3.2))
    img_elem = img_p._p
    img_p._p.getparent().remove(img_p._p)

    insert_before(vi_para, cap_elem)
    insert_before(vi_para, img_elem)

# ─── update Abstract ─────────────────────────────────────────────────────────
abstract_para = find("Abstract")
abstract_para.add_run(
    " The design is validated on a Cyclone IV E FPGA (EP4CE10F17C8): "
    "hardware synthesis confirms that round adds only 46 logic elements "
    "versus trunc at K=6; on-chip inference shows K=6 trunc misclassifies "
    "(argmax 1→3) while round maintains 96.8% agreement with exact inference "
    "across 1000 test samples."
).font.color.rgb = BLUE

# ─── update Conclusion ───────────────────────────────────────────────────────
last_concl = None
in_concl = False
for p in paras:
    if "VII. CONCLUSION" in p.text:
        in_concl = True
    if in_concl and p.text.strip():
        last_concl = p
    if in_concl and "REFERENCES" in p.text:
        break

if last_concl:
    last_concl.add_run(
        " Hardware validation on the EP4CE10 FPGA confirms these findings: "
        "K=6 truncation causes misclassification in on-chip inference "
        "(argmax shift 1→3), matching the Python simulation prediction, "
        "while deterministic rounding at K=6 retains 96.8% Top-1 agreement. "
        "Synthesis results show round incurs only +46 logic elements versus "
        "trunc, confirming the near-zero hardware overhead claimed in the analysis."
    ).font.color.rgb = BLUE

doc.save(DST)
print("Saved v3 report")
