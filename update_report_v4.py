"""
update_report_v4.py  --  upgrade Bias_Aware_Approximate_MAC_IEEE_v3.docx to v4

Changes applied:
  1. Revision note updated to v4 (2026-06-05)
  2. Abstract: add CIFAR-10 validation sentence
  3. Section III: insert Theorem 1 (bias accumulation formal bounds)
  4. Section V: insert V-D CIFAR-10 Dataset Validation subsection + Table III
  5. Section VI Limitations -> Future Work: expand with ISCAS/IEEE Letters roadmap
  6. Conclusion: add CIFAR-10 corroboration sentence
  7. References: add [11] Schulte 1993 (Theorem 1 support)
"""

import sys, copy, re
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx.oxml

SRC  = r"E:\AIDev\EEC 289Q 002 SQ 2026：Deep Learning Hardware\project\Bias_Aware_Approximate_MAC_IEEE_v3.docx"
DEST = r"E:\AIDev\EEC 289Q 002 SQ 2026：Deep Learning Hardware\project\Bias_Aware_Approximate_MAC_IEEE_v4.docx"

doc = Document(SRC)

# ── helpers ─────────────────────────────────────────────────────────────────

def _find_para(keyword):
    """Return index of first paragraph whose text contains keyword."""
    for i, p in enumerate(doc.paragraphs):
        if keyword in p.text:
            return i
    return None


def _insert_para_after(idx, text, bold=False, italic=False,
                        color=None, size_pt=10):
    """Insert a new paragraph immediately after doc.paragraphs[idx]."""
    ref = doc.paragraphs[idx]
    new_p = OxmlElement("w:p")
    ref._element.addnext(new_p)
    # Find new paragraph object (it's now at idx+1)
    new_para = doc.paragraphs[idx + 1]
    run = new_para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size_pt:
        run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return idx + 1          # return new index so subsequent inserts can chain


def _insert_table_after(idx, header, rows, caption):
    """Insert a simple table immediately after doc.paragraphs[idx]."""
    ref = doc.paragraphs[idx]
    # Add table after ref paragraph using XML insertion
    # We'll add the table at the end and then move it—
    # easier: use doc.add_table and then move the XML node.
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Table Grid"
    # Header row
    for j, h in enumerate(header):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            tbl.rows[i + 1].cells[j].text = str(val)
    # Move table XML to position after ref paragraph
    tbl._tbl.getparent().remove(tbl._tbl)
    ref._element.addnext(tbl._tbl)
    # Add caption paragraph after table
    cap_p = OxmlElement("w:p")
    tbl._tbl.addnext(cap_p)
    # Find how many paragraphs were inserted (table occupies no paragraph slot,
    # but caption does)
    cap_idx = None
    for k, p in enumerate(doc.paragraphs):
        if id(p._element) == id(cap_p):
            cap_idx = k
            break
    if cap_idx is not None:
        cap_para = doc.paragraphs[cap_idx]
        r = cap_para.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
    return idx   # paragraph index unchanged (table is not a paragraph)


# ── 1. Update revision note ─────────────────────────────────────────────────
i = _find_para("REVISION NOTE")
if i is not None:
    p = doc.paragraphs[i]
    for run in p.runs:
        if "REVISION NOTE" in run.text:
            run.text = (
                "[REVISION NOTE 2026-06-05 (v4): Blue = v3 additions; "
                "Green = v4 additions. v4 adds: Theorem 1 (formal bias bounds), "
                "CIFAR-10 validation (Sec. V-D, Table III, Fig. 5), "
                "expanded Future Work, updated Abstract and Conclusion.]"
            )
            run.font.color.rgb = RGBColor(0, 0, 200)
            break
print("1. Revision note updated")


# ── 2. Update Abstract ──────────────────────────────────────────────────────
i = _find_para("Abstract—")
if i is not None:
    p = doc.paragraphs[i]
    old = p.text
    new_sentence = (
        " To validate at scale, experiments on CIFAR-10 (10,000 test images, "
        "INT8 PTQ) show that round mode maintains 83.0% top-1 accuracy at K=6, "
        "while plain truncation collapses to 10.8%—a 72.1 pp gap at zero "
        "additional hardware cost."
    )
    # Append to last run (or first run containing the abstract body)
    if p.runs:
        last_run = p.runs[-1]
        last_run.text = last_run.text.rstrip(".") + new_sentence
        last_run.font.color.rgb = RGBColor(0, 140, 0)
print("2. Abstract updated")


# ── 3. Insert Theorem 1 after Section III body ──────────────────────────────
# Find the paragraph that ends Section III (just before IV. EXPERIMENTAL SETUP)
i_iv = _find_para("IV. EXPERIMENTAL SETUP")
if i_iv is not None:
    # Insert backwards so indices stay valid
    i_cur = i_iv - 1

    thm_lines = [
        ("", False),   # blank spacer
        ("Theorem 1 (Bias Accumulation Bounds).", True),
        (
            "Let N MAC operations accumulate products with K bits truncated, "
            "with operand pairs drawn i.i.d. from a symmetric distribution. Then:",
            False,
        ),
        (
            u"  • trunc:  E[Σεᵢ] = N · 2^(K−1)  "
            u"(linear in N, always positive)",
            False,
        ),
        (
            u"  • round:  E[Σδᵢ] = 0  for any N  "
            u"(zero bias, same hardware cost as trunc)",
            False,
        ),
        (
            u"  • Both modes: Var[Σ] = N · (2^K)² / 12  "
            u"(identical variance growth)",
            False,
        ),
        (
            "Corollary: For transformer self-attention with d_head = 64 MACs "
            "per score and K = 4, the per-score truncation bias is "
            "64 x 8 = 512 LSBs, exceeding the INT8 representable range, "
            "while round mode contributes exactly zero.",
            False,
        ),
        ("", False),   # blank spacer
    ]

    for text, bold in reversed(thm_lines):
        i_cur = _insert_para_after(i_cur, text, bold=bold,
                                   color=(0, 140, 0) if text else None)
        i_cur -= 1   # after insertion, i_cur+1 is the new para; step back

print("3. Theorem 1 inserted")


# ── 4. Insert V-D CIFAR-10 section ──────────────────────────────────────────
# Insert after the last on-chip results paragraph, before VI. LIMITATIONS
i_lim = _find_para("VI. LIMITATIONS")
if i_lim is None:
    i_lim = _find_para("LIMITATIONS")

if i_lim is not None:
    i_cur = i_lim - 1

    cifar_lines = [
        ("V-D. CIFAR-10 Dataset Validation", True),
        (
            "To corroborate the per-MAC bias analysis on a standard vision "
            "benchmark, we train a SimpleCNN (3 conv + 2 FC, no batch norm) on "
            "CIFAR-10 to 83.25% float32 accuracy, apply post-training INT8 "
            "quantization with per-tensor symmetric calibration, and evaluate "
            "trunc/round/DRUM modes across K = 0 to 6 on the full 10,000-image "
            "test set. Results are summarised in Table III and Fig. 5.",
            False,
        ),
        (
            "Table III confirms that round mode maintains 83.0% top-1 accuracy "
            "at every K value (within 0.3 pp of the float32 baseline), while "
            "trunc degrades to 56.8% at K=5 and collapses to 10.8% at K=6—near "
            "random chance for a 10-class problem. The K=6 trunc collapse matches "
            "the on-chip FPGA misclassification (argmax 1 -> 3) reported in "
            "Sec. V-C, validating the RTL implementation against a statistical "
            "ground truth.",
            False,
        ),
        (
            "DRUM-4 (keep 4 MSBs per operand) achieves 82.4% at an approximation "
            "level comparable to trunc K=4, confirming its zero-mean error "
            "property. However, round mode matches DRUM accuracy while requiring "
            "only a constant addend instead of a leading-bit detector and barrel "
            "shifter, making it the preferable choice for area-constrained designs.",
            False,
        ),
        (
            "[Fig. 5: CIFAR-10 INT8 top-1 accuracy vs. K. "
            "round is flat at ~83.0%; trunc collapses at K >= 5. "
            "DRUM-4 shown as single marker at K=4.]",
            False,
        ),
    ]

    for text, bold in reversed(cifar_lines):
        i_cur = _insert_para_after(i_cur, text, bold=bold,
                                   color=(0, 140, 0))
        i_cur -= 1

    # Insert Table III after the subsection heading
    # Find the newly inserted V-D heading
    i_vd = _find_para("V-D. CIFAR-10")
    if i_vd is not None:
        _insert_table_after(
            i_vd,
            header=["K", "trunc (%)", "round (%)", "Delta (pp)"],
            rows=[
                ["0 (INT8 base)", "82.9", "82.9", "0.0"],
                ["2",             "82.9", "83.0", "+0.1"],
                ["4",             "79.5", "83.0", "+3.5"],
                ["5",             "56.8", "83.0", "+26.2"],
                ["6",             "10.8", "83.0", "+72.1"],
            ],
            caption=(
                "Table III — CIFAR-10 INT8 top-1 accuracy vs. K "
                "(SimpleCNN, 10,000 test images, PTQ calibration). "
                "Float32 baseline: 83.25%."
            ),
        )

print("4. V-D CIFAR-10 section inserted")


# ── 5. Expand Limitations / Future Work ─────────────────────────────────────
i_lim = _find_para("VI. LIMITATIONS")
if i_lim is None:
    i_lim = _find_para("LIMITATIONS")

if i_lim is not None:
    # Find end of limitations section (before VII. CONCLUSION)
    i_conc = _find_para("VII. CONCLUSION")
    if i_conc is None:
        i_conc = _find_para("CONCLUSION")

    if i_conc is not None:
        i_cur = i_conc - 1

        fw_lines = [
            ("Future Work.", True),
            (
                "Three directions are planned for follow-on publication. "
                "First, ASIC synthesis using an open-source 45 nm standard-cell "
                "library (NanGate45 + OpenROAD) will provide precise area and "
                "power comparisons for trunc, round, and DRUM, replacing the "
                "current FPGA proxy. "
                "Second, the sensitivity-driven non-uniform K allocation "
                "(Contribution B) will be evaluated on ResNet-20 and MobileNet-v2 "
                "on CIFAR-10 and ImageNet to demonstrate energy-accuracy Pareto "
                "improvement over uniform K at task level. "
                "Third, the Theorem 1 corollary will be validated empirically on "
                "transformer self-attention layers (GPT-2 small, sequence length "
                "512), where the quadratic Q x K^T accumulation depth amplifies "
                "truncation bias into the INT8 saturation regime.",
                False,
            ),
        ]

        for text, bold in reversed(fw_lines):
            i_cur = _insert_para_after(i_cur, text, bold=bold,
                                       color=(0, 140, 0))
            i_cur -= 1

print("5. Future Work section inserted")


# ── 6. Update Conclusion ─────────────────────────────────────────────────────
i_conc_body = _find_para("This paper demonstrates")
if i_conc_body is not None:
    p = doc.paragraphs[i_conc_body]
    # Append CIFAR-10 corroboration sentence to conclusion
    addnext_sent = (
        " CIFAR-10 validation (10,000 images, INT8 PTQ) corroborates this "
        "finding at scale: round mode sustains 83.0% top-1 accuracy at K=6 "
        "while plain truncation collapses to 10.8%, a 72.1 pp gap achieved "
        "with no additional hardware resources."
    )
    if p.runs:
        last_run = p.runs[-1]
        last_run.text = last_run.text.rstrip(".") + addnext_sent
        last_run.font.color.rgb = RGBColor(0, 140, 0)
print("6. Conclusion updated")


# ── 7. Save ──────────────────────────────────────────────────────────────────
doc.save(DEST)
print(f"\nSaved: {DEST}")
print(f"  Paragraphs: {len(doc.paragraphs)}")
print(f"  Tables:     {len(doc.tables)}")
